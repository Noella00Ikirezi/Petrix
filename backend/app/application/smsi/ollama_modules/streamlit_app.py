#!/usr/bin/env python3
"""
SMSI-Ollama : Interface Web Streamlit
======================================
Interface graphique locale pour la suite d'outils SMSI
"""

import streamlit as st
import sys
from pathlib import Path

# Ajouter le répertoire parent au path
sys.path.insert(0, str(Path(__file__).parent))

# Imports pour les formats Office
import pandas as pd
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False

from config.settings import ORGANIZATION, OLLAMA_MODEL
from modules.ollama_client import OllamaClient
from modules.doc_generator import DocumentGenerator
from modules.compliance_analyzer import ComplianceAnalyzer
from modules.smsi_assistant import SMSIAssistant

# Configuration de la page
st.set_page_config(
    page_title="SMSI-Ollama",
    page_icon="🔐",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS personnalisé
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: #1E88E5;
        text-align: center;
        margin-bottom: 1rem;
    }
    .status-ok { color: #4CAF50; }
    .status-error { color: #F44336; }
    .info-box {
        background-color: #E3F2FD;
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)


def extract_text_from_file(uploaded_file):
    """Extrait le texte d'un fichier uploadé selon son format"""
    filename = uploaded_file.name.lower()
    content = ""

    try:
        if filename.endswith(('.txt', '.md')):
            # Fichiers texte
            content = uploaded_file.read().decode('utf-8')

        elif filename.endswith('.csv'):
            # CSV
            df = pd.read_csv(uploaded_file)
            content = df.to_string()

        elif filename.endswith('.xlsx'):
            # Excel moderne
            if not XLSX_AVAILABLE:
                raise ImportError("openpyxl non installé")

            # Lire toutes les feuilles
            excel_file = pd.ExcelFile(uploaded_file)
            all_sheets = []
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                all_sheets.append(f"\n=== Feuille: {sheet_name} ===\n")
                all_sheets.append(df.to_string())
            content = "\n".join(all_sheets)

        elif filename.endswith('.xls'):
            # Excel ancien format
            df = pd.read_excel(uploaded_file, engine='xlrd')
            content = df.to_string()

        elif filename.endswith('.docx'):
            # Word moderne
            if not DOCX_AVAILABLE:
                raise ImportError("python-docx non installé")

            doc = DocxDocument(uploaded_file)
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            # Aussi extraire les tableaux
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    paragraphs.append(row_text)
            content = "\n".join(paragraphs)

        else:
            raise ValueError(f"Format non supporté: {filename}")

    except Exception as e:
        raise Exception(f"Erreur lecture {filename}: {str(e)}")

    return content


def check_ollama_connection():
    """Vérifie la connexion à Ollama"""
    try:
        client = OllamaClient()
        models = client.list_models()
        return True, models
    except Exception as e:
        return False, str(e)


def init_session_state():
    """Initialise l'état de session"""
    if "messages" not in st.session_state:
        st.session_state.messages = []
    if "assistant" not in st.session_state:
        st.session_state.assistant = None


def sidebar():
    """Barre latérale avec statut et configuration"""
    with st.sidebar:
        st.image("https://img.icons8.com/color/96/000000/security-checked.png", width=80)
        st.title("SMSI-Ollama")

        # Statut Ollama
        st.subheader("📡 Statut")
        connected, result = check_ollama_connection()

        if connected:
            st.success(f"✓ Ollama connecté")
            st.caption(f"Modèles : {', '.join(result[:3])}")
        else:
            st.error(f"✗ Ollama non disponible")
            st.caption("Lancez : ollama serve")

        st.divider()

        # Configuration
        st.subheader("⚙️ Configuration")
        st.text(f"Organisation : {ORGANIZATION.get('name', 'N/A')}")
        st.text(f"Modèle : {OLLAMA_MODEL}")
        st.text(f"Référentiels : {len(ORGANIZATION.get('frameworks', []))}")

        st.divider()

        # Navigation
        st.subheader("📋 Navigation")
        page = st.radio(
            "Choisir un outil :",
            ["🏠 Accueil", "💬 Assistant SMSI", "📄 Générateur", "🔍 Analyseur"],
            label_visibility="collapsed"
        )

        return page, connected


def page_home():
    """Page d'accueil"""
    st.markdown('<p class="main-header">🔐 SMSI-Ollama</p>', unsafe_allow_html=True)
    st.markdown("<center>Suite d'outils IA pour la gestion du Système de Management de la Sécurité de l'Information</center>", unsafe_allow_html=True)

    st.divider()

    # Cartes des fonctionnalités
    col1, col2 = st.columns(2)

    with col1:
        st.markdown("### 💬 Assistant SMSI")
        st.markdown("""
        Posez vos questions sur :
        - ISO 27001:2022
        - RGPD et conformité
        - Bonnes pratiques sécurité
        - Implémentation SMSI
        """)

        st.markdown("### 📄 Générateur de Documents")
        st.markdown("""
        Générez automatiquement :
        - Politiques de sécurité
        - Procédures opérationnelles
        - Checklists d'audit
        """)

    with col2:
        st.markdown("### 🔍 Analyseur de Conformité")
        st.markdown("""
        Analysez vos documents :
        - Vérification des exigences
        - Détection des lacunes
        - Recommandations
        """)

        st.markdown("### 📊 Configuration")
        st.markdown(f"""
        **Organisation :** {ORGANIZATION.get('name')}
        **Secteur :** {ORGANIZATION.get('sector')}
        **Référentiels :** {', '.join(ORGANIZATION.get('frameworks', []))}
        """)


def page_assistant():
    """Page Assistant SMSI interactif"""
    st.header("💬 Assistant SMSI")
    st.caption("Posez vos questions sur ISO 27001, RGPD, et la sécurité de l'information")

    # Initialiser l'assistant
    if st.session_state.assistant is None:
        try:
            st.session_state.assistant = SMSIAssistant()
        except Exception as e:
            st.error(f"Erreur d'initialisation : {e}")
            return

    # Afficher l'historique des messages
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    # Zone de saisie
    if prompt := st.chat_input("Posez votre question..."):
        # Ajouter le message utilisateur
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)

        # Générer la réponse
        with st.chat_message("assistant"):
            with st.spinner("Réflexion en cours..."):
                try:
                    response = st.session_state.assistant.ask(prompt)
                    st.markdown(response)
                    st.session_state.messages.append({"role": "assistant", "content": response})
                except Exception as e:
                    st.error(f"Erreur : {e}")

    # Bouton pour effacer l'historique
    if st.button("🗑️ Effacer la conversation"):
        st.session_state.messages = []
        st.session_state.assistant = SMSIAssistant()
        st.rerun()


def page_generator():
    """Page Générateur de documents"""
    st.header("📄 Générateur de Documents SMSI")
    st.caption("Générez des politiques, procédures et checklists conformes ISO 27001")

    col1, col2 = st.columns([1, 1])

    with col1:
        doc_type = st.selectbox(
            "Type de document",
            ["POL - Politique", "PRO - Procédure", "CHK - Checklist"],
            help="Choisissez le type de document à générer"
        )

        ref_iso = st.text_input(
            "Référence ISO",
            placeholder="Ex: A.9.1.1",
            help="Numéro de la mesure ISO 27001"
        )

        measure_name = st.text_input(
            "Nom de la mesure",
            placeholder="Ex: Politique de contrôle d'accès",
            help="Description de la mesure"
        )

    with col2:
        context = st.text_area(
            "Contexte additionnel (optionnel)",
            placeholder="Informations spécifiques à votre organisation...",
            height=150
        )

    if st.button("🚀 Générer le document", type="primary"):
        if not ref_iso or not measure_name:
            st.warning("Veuillez remplir la référence ISO et le nom de la mesure")
            return

        with st.spinner("Génération en cours... Cela peut prendre quelques instants."):
            try:
                generator = DocumentGenerator()
                doc_type_code = doc_type.split(" - ")[0]

                content = generator.generate(
                    ref_iso=ref_iso,
                    measure_name=measure_name,
                    doc_type=doc_type_code,
                    additional_context=context or ""
                )

                st.success("Document généré avec succès !")

                # Afficher le document
                st.markdown("### Document généré")
                st.markdown(content)

                # Bouton de téléchargement
                st.download_button(
                    label="📥 Télécharger (Markdown)",
                    data=content,
                    file_name=f"{doc_type_code}_{ref_iso.replace('.', '_')}.md",
                    mime="text/markdown"
                )

            except Exception as e:
                st.error(f"Erreur lors de la génération : {e}")


def page_analyzer():
    """Page Analyseur de conformité"""
    st.header("🔍 Analyseur de Conformité")
    st.caption("Analysez vos documents par rapport aux exigences ISO 27001")

    # Formats supportés
    supported_formats = ["txt", "md", "csv", "xlsx", "xls", "docx"]

    # Upload de fichier
    uploaded_file = st.file_uploader(
        "Chargez un document à analyser",
        type=supported_formats,
        help="Formats supportés : TXT, Markdown, CSV, Excel (.xlsx, .xls), Word (.docx)"
    )

    # Afficher les formats supportés
    st.info("📁 **Formats supportés :** TXT, Markdown, CSV, Excel (.xlsx, .xls), Word (.docx)")

    # Sélection des exigences
    requirements = st.multiselect(
        "Exigences à vérifier",
        [
            "A.5.1 - Politiques de sécurité",
            "A.5.2 - Rôles et responsabilités",
            "A.9.1.1 - Politique de contrôle d'accès",
            "A.9.2.1 - Enregistrement utilisateurs",
            "A.9.2.3 - Gestion des droits privilégiés",
            "A.12.4.1 - Journalisation",
            "A.16.1.1 - Gestion des incidents",
            "A.18.1.1 - Législation applicable",
            "A.18.1.3 - Protection des enregistrements",
        ],
        default=["A.5.1 - Politiques de sécurité"]
    )

    # Option pour voir le contenu extrait
    show_extracted = st.checkbox("Afficher le contenu extrait", value=False)

    if st.button("🔍 Analyser", type="primary"):
        if not uploaded_file:
            st.warning("Veuillez charger un document")
            return

        if not requirements:
            st.warning("Veuillez sélectionner au moins une exigence")
            return

        with st.spinner("Extraction du contenu..."):
            try:
                # Extraire le contenu selon le format
                content = extract_text_from_file(uploaded_file)

                if not content.strip():
                    st.error("Le document semble vide ou le contenu n'a pas pu être extrait.")
                    return

                # Afficher le contenu extrait si demandé
                if show_extracted:
                    with st.expander("📄 Contenu extrait", expanded=False):
                        st.text(content[:5000] + ("..." if len(content) > 5000 else ""))
                        st.caption(f"Longueur totale : {len(content)} caractères")

            except Exception as e:
                st.error(f"Erreur lors de l'extraction : {e}")
                return

        with st.spinner("Analyse IA en cours..."):
            try:
                # Extraire les codes d'exigences
                req_codes = [r.split(" - ")[0] for r in requirements]

                analyzer = ComplianceAnalyzer()
                report = analyzer.analyze_document(
                    document_content=content,
                    requirements=req_codes,
                    document_name=uploaded_file.name
                )

                st.success("Analyse terminée !")

                # Afficher le rapport
                st.markdown("### Rapport d'analyse")
                formatted_report = analyzer.format_report_markdown(report)
                st.markdown(formatted_report)

                # Téléchargement
                st.download_button(
                    label="📥 Télécharger le rapport",
                    data=formatted_report,
                    file_name=f"analyse_{uploaded_file.name}.md",
                    mime="text/markdown"
                )

            except Exception as e:
                st.error(f"Erreur lors de l'analyse : {e}")


def main():
    """Point d'entrée principal"""
    init_session_state()

    # Sidebar et navigation
    page, connected = sidebar()

    if not connected:
        st.warning("⚠️ Ollama n'est pas accessible. Lancez `ollama serve` dans un terminal.")

    # Routage des pages
    if page == "🏠 Accueil":
        page_home()
    elif page == "💬 Assistant SMSI":
        page_assistant()
    elif page == "📄 Générateur":
        page_generator()
    elif page == "🔍 Analyseur":
        page_analyzer()


if __name__ == "__main__":
    main()
